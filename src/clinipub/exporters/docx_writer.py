import os
import pandas as pd
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


class JournalDocxExporter:
    """Generates a Microsoft Word document containing a native table styled
    to match the specific formatting requirements of medical journals.
    """

    STYLES = {
        "nejm": {
            "font": "Times New Roman",
            "size": 10,
            "top_border": {"val": "single", "sz": "12", "color": "000000"},
            "bottom_border": {"val": "single", "sz": "12", "color": "000000"},
            "header_bottom": {"val": "single", "sz": "6", "color": "000000"},
            "inside_h": {"val": "none", "sz": "0", "color": "auto"},
        },
        "jama": {
            "font": "Arial",
            "size": 9.5,
            "top_border": {"val": "single", "sz": "18", "color": "111111"},
            "bottom_border": {"val": "single", "sz": "18", "color": "111111"},
            "header_bottom": {"val": "single", "sz": "6", "color": "111111"},
            "inside_h": {"val": "single", "sz": "4", "color": "E2E8F0"},
        },
        "lancet": {
            "font": "Arial",
            "size": 9,
            "top_border": {"val": "single", "sz": "6", "color": "333333"},
            "bottom_border": {"val": "single", "sz": "6", "color": "333333"},
            "header_bottom": {"val": "single", "sz": "6", "color": "333333"},
            "inside_h": {"val": "single", "sz": "4", "color": "EEEEEE"},
        },
    }

    def __init__(self, table_df: pd.DataFrame, journal: str = "nejm"):
        if table_df.empty:
            raise ValueError("The input DataFrame cannot be empty.")
            
        journal_key = journal.lower().strip()
        if journal_key not in self.STYLES:
            raise ValueError(f"Unsupported journal style '{journal}'. Choose from: {list(self.STYLES.keys())}")
            
        self.table_df = table_df.copy()
        self.journal = journal_key
        self.cfg = self.STYLES[self.journal]

    def _set_cell_margins(self, cell, top=100, bottom=100, start=150, end=150):
        """Injects explicit padding values directly into the cell's XML structures."""
        tcPr = cell._tc.get_or_add_tcPr()
        tcMar = OxmlElement('w:tcMar')
        for m, val in [('top', top), ('bottom', bottom), ('left', start), ('right', end)]:
            node = OxmlElement(f'w:{m}')
            node.set(qn('w:w'), str(val))
            node.set(qn('w:type'), 'dxa')
            tcMar.append(node)
        tcPr.append(tcMar)

    def _apply_table_borders(self, table):
        """Modifies table properties using OpenXML elements to enforce border rules."""
        tblPr = table._tbl.tblPr
        tblBorders = OxmlElement('w:tblBorders')
        
        # Configure the structural layout borders based on style configurations
        for border_name, cfg in [('top', self.cfg['top_border']), 
                                 ('bottom', self.cfg['bottom_border']),
                                 ('insideH', self.cfg['inside_h'])]:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), cfg['val'])
            border.set(qn('w:sz'), cfg['sz'])
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), cfg['color'])
            tblBorders.append(border)
            
        # Eliminate messy vertical borders completely
        for side in ['left', 'right', 'insideV']:
            border = OxmlElement(f'w:{side}')
            border.set(qn('w:val'), 'none')
            tblBorders.append(border)
            
        tblPr.append(tblBorders)

    def save(self, filepath: str):
        """Constructs a native OpenXML table inside a fresh Word Document and saves it."""
        doc = Document()
        
        # Set section to landscape to support wide multi-arm baseline tables
        section = doc.sections[-1]
        section.orientation = Inches(11) > Inches(8.5)
        section.page_width = Inches(11)
        section.page_height = Inches(8.5)

        # Build clean column structures
        headers = [str(col) for col in self.table_df.columns]
        has_index = self.table_df.index.name is not None or list(self.table_df.index)
        
        if has_index:
            headers = ["Variable"] + headers

        table = doc.add_table(rows=1, cols=len(headers))
        self._apply_table_borders(table)

        # 1. Format Header Row
        hdr_cells = table.rows[0].cells
        for idx, text in enumerate(headers):
            cell = hdr_cells[idx]
            cell.text = text
            self._set_cell_margins(cell, top=140, bottom=140)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.font.name = self.cfg['font']
                run.font.size = Pt(self.cfg['size'])
                run.font.bold = True

        # Apply bottom line rule specifically underneath the header row block
        trPr = table.rows[0]._tr.get_or_add_trPr()
        tblHeader = OxmlElement('w:tblHeader')
        trPr.append(tblHeader)

        # 2. Populate Data Rows
        for row_idx, idx_val in enumerate(self.table_df.index):
            row_cells = table.add_row().cells
            
            # Map structural columns out
            row_data = []
            if has_index:
                row_data.append(str(idx_val))
            row_data.extend([str(val) for val in self.table_df.iloc[row_idx]])

            for col_idx, cell_value in enumerate(row_data):
                cell = row_cells[col_idx]
                cell.text = cell_value
                
                # Check for sub-category flags to enforce nested indentation rules
                is_sub_category = cell_value.startswith("  ") or cell_value.startswith("\t")
                left_pad = 300 if (col_idx == 0 and is_sub_category) else 150
                self._set_cell_margins(cell, top=80, bottom=80, start=left_pad)
                
                p = cell.paragraphs[0]
                # Left align row headers; right align quantitative metrics
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT if col_idx == 0 else WD_ALIGN_PARAGRAPH.RIGHT
                
                for run in p.runs:
                    run.font.name = self.cfg['font']
                    run.font.size = Pt(self.cfg['size'])
                    # Bold main section categories if they contain no data
                    if col_idx == 0 and not is_sub_category and str(self.table_df.iloc[row_idx, 0]) == "":
                        run.font.bold = True

        doc.save(filepath)
